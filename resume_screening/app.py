import streamlit as st
import os
import PyPDF2
from docx import Document
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
from collections import Counter
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
import re

# Page configuration with modern theme
st.set_page_config(
    page_title="Resume Screening & Keyword Extraction",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern UI
st.markdown("""
    <style>
    .main {
        padding-top: 2rem;
    }
    .stMetric {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .keyword-card {
        background-color: #ffffff;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        border: 1px solid #e0e0e0;
        margin: 0.25rem;
        display: inline-block;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .extract-box {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-radius: 0.5rem;
        border: 1px solid #dee2e6;
        margin: 1rem 0;
    }
    h1 {
        color: #1f77b4;
        border-bottom: 3px solid #1f77b4;
        padding-bottom: 0.5rem;
    }
    h2 {
        color: #2c3e50;
        margin-top: 2rem;
    }
    .stProgress > div > div > div > div {
        background-color: #1f77b4;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize spaCy model
@st.cache_resource
def load_nlp_model():
    try:
        return spacy.load('en_core_web_sm')
    except OSError:
        st.error("spaCy model 'en_core_web_sm' not found. Please install it using: python -m spacy download en_core_web_sm")
        return None

nlp = load_nlp_model()

SKILLS = ["FullStack Developer", "Python", "JavaScript", "Java", "React", "Node.js", 
          "Django", "Flask", "SQL", "MongoDB", "Machine Learning", "Data Analysis", "NLP"]

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    text = ""
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_uploaded_pdf(uploaded_file):
    """Extract text from uploaded PDF file"""
    text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(uploaded_file.read()))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
    return text

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_uploaded_docx(uploaded_file):
    """Extract text from uploaded DOCX file"""
    try:
        doc = Document(BytesIO(uploaded_file.read()))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def preprocess_text(text):
    """Preprocess text using spaCy"""
    if nlp is None:
        return text.lower()
    doc = nlp(text.lower())
    tokens = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
    return " ".join(tokens)

def extract_skills(text):
    """Extract skills from text"""
    found_skills = []
    text_lower = text.lower()
    for skill in SKILLS:
        if skill.lower() in text_lower:
            found_skills.append(skill)
    return found_skills

def extract_keywords_advanced(text, top_n=20):
    """Extract keywords using TF-IDF and NLP"""
    if not text or len(text.strip()) < 10:
        return []
    
    if nlp is None:
        # Fallback to simple extraction
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        word_freq = Counter(words)
        return [word for word, _ in word_freq.most_common(top_n)]
    
    # Extract nouns and important terms using spaCy
    doc = nlp(text)
    
    # Extract keywords (nouns, proper nouns, and important adjectives)
    keywords = []
    for token in doc:
        if (token.pos_ in ['NOUN', 'PROPN'] and 
            not token.is_stop and 
            len(token.text) > 2 and
            token.is_alpha):
            keywords.append(token.lemma_.lower())
    
    # Also extract named entities
    entities = [ent.text for ent in doc.ents if ent.label_ in ['ORG', 'PRODUCT', 'TECHNOLOGY']]
    keywords.extend([e.lower() for e in entities])
    
    # Use TF-IDF for additional keyword extraction
    try:
        vectorizer = TfidfVectorizer(max_features=top_n, stop_words='english', ngram_range=(1, 2))
        tfidf_matrix = vectorizer.fit_transform([text])
        feature_names = vectorizer.get_feature_names_out()
        tfidf_scores = tfidf_matrix.toarray()[0]
        
        # Combine NLP and TF-IDF keywords
        tfidf_keywords = [feature_names[i] for i in tfidf_scores.argsort()[-top_n:][::-1] if tfidf_scores[i] > 0]
        
        # Merge and deduplicate
        all_keywords = list(set(keywords + tfidf_keywords))
        keyword_freq = Counter(all_keywords)
        
        return [word for word, _ in keyword_freq.most_common(top_n)]
    except:
        keyword_freq = Counter(keywords)
        return [word for word, _ in keyword_freq.most_common(top_n)]

def extract_entities(text):
    """Extract named entities from text"""
    if nlp is None:
        return {}
    
    doc = nlp(text)
    entities = {}
    
    for ent in doc.ents:
        if ent.label_ not in entities:
            entities[ent.label_] = []
        entities[ent.label_].append(ent.text)
    
    # Deduplicate
    for label in entities:
        entities[label] = list(set(entities[label]))
    
    return entities

def extract_contact_info(text):
    """Extract contact information from text"""
    contact_info = {
        'email': [],
        'phone': [],
        'linkedin': [],
        'github': []
    }
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    contact_info['email'] = list(set(emails))
    
    # Phone pattern
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    contact_info['phone'] = list(set([''.join(p) for p in phones if any(p)]))
    
    # LinkedIn pattern
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    contact_info['linkedin'] = list(set(linkedin))
    
    # GitHub pattern
    github_pattern = r'github\.com/[\w-]+'
    github = re.findall(github_pattern, text, re.IGNORECASE)
    contact_info['github'] = list(set(github))
    
    return contact_info

def load_resume_text(file_path):
    """Load resume text from file path"""
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format")

def main():
    st.title("📄 Resume Screening & Keyword Extraction")
    st.markdown("---")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        resume_folder = st.text_input(
            "Resume Folder Path",
            value=os.path.join(os.getcwd(), "resumes"),
            help="Path to folder containing resume files"
        )
        
        st.markdown("---")
        st.header("📤 Upload Resume")
        uploaded_file = st.file_uploader(
            "Upload a resume file",
            type=['pdf', 'docx'],
            help="Upload a PDF or DOCX resume file for analysis"
        )
        
        st.markdown("---")
        st.header("📊 Settings")
        top_keywords = st.slider("Number of Keywords to Extract", 5, 50, 20)
        show_entities = st.checkbox("Show Named Entities", value=True)
        show_contact = st.checkbox("Show Contact Information", value=True)
    
    # Main content area
    if uploaded_file is not None:
        # Process uploaded file
        st.header("📋 Resume Analysis")
        
        with st.spinner("Processing resume..."):
            if uploaded_file.name.endswith('.pdf'):
                raw_text = extract_text_from_uploaded_pdf(uploaded_file)
            elif uploaded_file.name.endswith('.docx'):
                raw_text = extract_text_from_uploaded_docx(uploaded_file)
            else:
                st.error("Unsupported file format")
                return
        
        if raw_text:
            # Display file info
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("File Name", uploaded_file.name)
            with col2:
                st.metric("Text Length", f"{len(raw_text):,} characters")
            with col3:
                st.metric("Word Count", f"{len(raw_text.split()):,} words")
            
            st.markdown("---")
            
            # Extract and display details
            tab1, tab2, tab3, tab4 = st.tabs(["📝 Extracted Text", "🔑 Keywords", "📊 Skills & Entities", "📧 Contact Info"])
            
            with tab1:
                st.subheader("Full Extracted Text")
                st.markdown(f'<div class="extract-box">{raw_text[:5000]}{"..." if len(raw_text) > 5000 else ""}</div>', 
                           unsafe_allow_html=True)
                if len(raw_text) > 5000:
                    with st.expander("View Full Text"):
                        st.text(raw_text)
            
            with tab2:
                st.subheader(f"Top {top_keywords} Keywords")
                keywords = extract_keywords_advanced(raw_text, top_n=top_keywords)
                
                if keywords:
                    # Display keywords as cards
                    keyword_html = '<div style="display: flex; flex-wrap: wrap; gap: 0.5rem; margin: 1rem 0;">'
                    for i, keyword in enumerate(keywords):
                        keyword_html += f'<div class="keyword-card">{keyword}</div>'
                    keyword_html += '</div>'
                    st.markdown(keyword_html, unsafe_allow_html=True)
                    
                    # Keyword frequency chart
                    keyword_freq = Counter([w.lower() for w in raw_text.split() if w.lower() in [k.lower() for k in keywords]])
                    if keyword_freq:
                        df_keywords = pd.DataFrame({
                            'Keyword': list(keyword_freq.keys())[:15],
                            'Frequency': list(keyword_freq.values())[:15]
                        })
                        fig = px.bar(df_keywords, x='Frequency', y='Keyword', 
                                     orientation='h', 
                                     title="Keyword Frequency",
                                     color='Frequency',
                                     color_continuous_scale='Blues')
                        fig.update_layout(height=400)
                        st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("No keywords extracted. The text might be too short or contain no meaningful content.")
            
            with tab3:
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("Extracted Skills")
                    skills = extract_skills(raw_text)
                    if skills:
                        for skill in skills:
                            st.markdown(f"- ✅ **{skill}**")
                    else:
                        st.info("No predefined skills found in the resume.")
                    
                    # Skills matching visualization
                    if skills:
                        st.markdown("---")
                        st.subheader("Skills Match Score")
                        match_score = (len(skills) / len(SKILLS)) * 100
                        st.progress(match_score / 100)
                        st.metric("Match Percentage", f"{match_score:.1f}%")
                
                with col2:
                    if show_entities and nlp is not None:
                        st.subheader("Named Entities")
                        entities = extract_entities(raw_text)
                        if entities:
                            for label, values in entities.items():
                                with st.expander(f"{label} ({len(values)})"):
                                    for value in values[:10]:  # Limit to first 10
                                        st.markdown(f"- {value}")
                        else:
                            st.info("No named entities found.")
            
            with tab4:
                if show_contact:
                    st.subheader("Contact Information")
                    contact_info = extract_contact_info(raw_text)
                    
                    if contact_info['email']:
                        st.markdown("**📧 Email Addresses:**")
                        for email in contact_info['email']:
                            st.markdown(f"- {email}")
                    
                    if contact_info['phone']:
                        st.markdown("**📞 Phone Numbers:**")
                        for phone in contact_info['phone']:
                            st.markdown(f"- {phone}")
                    
                    if contact_info['linkedin']:
                        st.markdown("**💼 LinkedIn:**")
                        for linkedin in contact_info['linkedin']:
                            st.markdown(f"- https://www.{linkedin}")
                    
                    if contact_info['github']:
                        st.markdown("**💻 GitHub:**")
                        for github in contact_info['github']:
                            st.markdown(f"- https://www.{github}")
                    
                    if not any(contact_info.values()):
                        st.info("No contact information found in the resume.")
                else:
                    st.info("Contact information extraction is disabled.")
    
    else:
        # Process resumes from folder
        st.header("📁 Batch Resume Processing")
        
        if not os.path.exists(resume_folder):
            st.warning(f"Resume folder not found: {resume_folder}")
            if st.button("Create Folder"):
                os.makedirs(resume_folder, exist_ok=True)
                st.success(f"Created folder: {resume_folder}")
                st.rerun()
        else:
            resume_files = [f for f in os.listdir(resume_folder) 
                          if f.endswith('.pdf') or f.endswith('.docx')]
            
            if not resume_files:
                st.info(f"No resume files found in: {resume_folder}")
            else:
                st.success(f"Found {len(resume_files)} resume file(s)")
                
                # Job description input
                st.subheader("📝 Job Description")
                job_description = st.text_area(
                    "Enter job description for similarity matching",
                    height=150,
                    value="""We are looking for a Python developer with experience in machine learning,
data analysis, and web frameworks like Django or Flask.
Must have skills in SQL and NLP."""
                )
                
                if st.button("🚀 Process Resumes", type="primary"):
                    with st.spinner("Processing resumes..."):
                        resume_texts = []
                        resume_data = []
                        
                        progress_bar = st.progress(0)
                        for idx, filename in enumerate(resume_files):
                            file_path = os.path.join(resume_folder, filename)
                            try:
                                raw_text = load_resume_text(file_path)
                                clean_text = preprocess_text(raw_text)
                                resume_texts.append(clean_text)
                                
                                # Extract details
                                keywords = extract_keywords_advanced(raw_text, top_n=top_keywords)
                                skills = extract_skills(raw_text)
                                entities_dict = extract_entities(raw_text) if show_entities else {}
                                contact = extract_contact_info(raw_text) if show_contact else {}
                                
                                resume_data.append({
                                    'filename': filename,
                                    'raw_text': raw_text,
                                    'keywords': keywords,
                                    'skills': skills,
                                    'entities': entities_dict,
                                    'contact': contact
                                })
                                
                                progress_bar.progress((idx + 1) / len(resume_files))
                            except Exception as e:
                                st.error(f"Error processing {filename}: {str(e)}")
                        
                        if resume_texts:
                            # Calculate similarity scores
                            job_description_clean = preprocess_text(job_description)
                            documents = resume_texts + [job_description_clean]
                            
                            try:
                                vectorizer = TfidfVectorizer()
                                tfidf_matrix = vectorizer.fit_transform(documents)
                                
                                job_vector = tfidf_matrix[-1]
                                resume_vectors = tfidf_matrix[:-1]
                                similarities = cosine_similarity(resume_vectors, job_vector).flatten()
                                
                                # Rank resumes
                                ranked_data = []
                                for idx, (filename, score) in enumerate(zip(resume_files, similarities)):
                                    resume_info = next((r for r in resume_data if r['filename'] == filename), None)
                                    if resume_info:
                                        ranked_data.append({
                                            **resume_info,
                                            'similarity_score': score
                                        })
                                
                                ranked_data = sorted(ranked_data, key=lambda x: x['similarity_score'], reverse=True)
                                
                                # Display results
                                st.markdown("---")
                                st.header("📊 Results Dashboard")
                                
                                # Summary metrics
                                col1, col2, col3, col4 = st.columns(4)
                                with col1:
                                    st.metric("Total Resumes", len(ranked_data))
                                with col2:
                                    avg_score = sum(r['similarity_score'] for r in ranked_data) / len(ranked_data) if ranked_data else 0
                                    st.metric("Average Score", f"{avg_score:.3f}")
                                with col3:
                                    max_score = max(r['similarity_score'] for r in ranked_data) if ranked_data else 0
                                    st.metric("Highest Score", f"{max_score:.3f}")
                                with col4:
                                    total_skills = sum(len(r['skills']) for r in ranked_data)
                                    st.metric("Total Skills Found", total_skills)
                                
                                # Similarity scores chart
                                df_scores = pd.DataFrame({
                                    'Resume': [r['filename'] for r in ranked_data],
                                    'Similarity Score': [r['similarity_score'] for r in ranked_data]
                                })
                                
                                fig = px.bar(df_scores, x='Resume', y='Similarity Score',
                                           title="Resume Similarity Scores",
                                           color='Similarity Score',
                                           color_continuous_scale='Viridis')
                                fig.update_xaxes(tickangle=-45)
                                fig.update_layout(height=400)
                                st.plotly_chart(fig, use_container_width=True)
                                
                                # Detailed results
                                st.markdown("---")
                                st.header("📋 Detailed Results")
                                
                                for rank, resume_info in enumerate(ranked_data, 1):
                                    with st.expander(f"#{rank} {resume_info['filename']} - Score: {resume_info['similarity_score']:.4f}", 
                                                   expanded=(rank == 1)):
                                        col1, col2 = st.columns(2)
                                        
                                        with col1:
                                            st.markdown("**🔑 Top Keywords:**")
                                            if resume_info['keywords']:
                                                keywords_display = " ".join([
                                                    f"`{kw}`" for kw in resume_info['keywords'][:15]
                                                ])
                                                st.markdown(keywords_display)
                                            else:
                                                st.info("No keywords extracted")
                                            
                                            st.markdown("**✅ Skills Found:**")
                                            if resume_info['skills']:
                                                for skill in resume_info['skills']:
                                                    st.markdown(f"- {skill}")
                                            else:
                                                st.info("No predefined skills found")
                                        
                                        with col2:
                                            if show_contact and resume_info['contact']:
                                                st.markdown("**📧 Contact Info:**")
                                                contact = resume_info['contact']
                                                if contact.get('email'):
                                                    st.markdown(f"Email: {', '.join(contact['email'][:2])}")
                                                if contact.get('phone'):
                                                    st.markdown(f"Phone: {', '.join(contact['phone'][:2])}")
                                            
                                            if show_entities and resume_info['entities']:
                                                st.markdown("**🏷️ Named Entities:**")
                                                for label, values in list(resume_info['entities'].items())[:3]:
                                                    st.markdown(f"**{label}:** {', '.join(values[:3])}")
                                        
                                        # Text preview
                                        text_preview = resume_info['raw_text'][:500]
                                        st.markdown(f"**Text Preview:**")
                                        st.text(text_preview + "..." if len(resume_info['raw_text']) > 500 else text_preview)
                                
                                progress_bar.empty()
                                
                            except Exception as e:
                                st.error(f"Error calculating similarity: {str(e)}")

if __name__ == "__main__":
    main()

