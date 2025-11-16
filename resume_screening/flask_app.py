from flask import Flask, render_template, request, redirect, url_for, flash
import os
import PyPDF2
from docx import Document
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from collections import Counter
import re
from io import BytesIO

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")

# Load spaCy model
try:
    nlp = spacy.load('en_core_web_sm')
except Exception:
    nlp = None

SKILLS = [
    "FullStack Developer", "Python", "JavaScript", "Java", "React", "Node.js",
    "Django", "Flask", "SQL", "MongoDB", "Machine Learning", "Data Analysis", "NLP"
]


def extract_text_from_pdf_filelike(file_like):
    text = ""
    reader = PyPDF2.PdfReader(file_like)
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def extract_text_from_pdf_path(pdf_path):
    with open(pdf_path, "rb") as f:
        return extract_text_from_pdf_filelike(f)


def extract_text_from_docx_path(docx_path):
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs])


def extract_text_from_docx_filelike(file_like):
    doc = Document(file_like)
    return "\n".join([p.text for p in doc.paragraphs])


def preprocess_text(text: str) -> str:
    if not text:
        return ""
    if nlp is None:
        return text.lower()
    doc = nlp(text.lower())
    tokens = [t.lemma_ for t in doc if not t.is_stop and t.is_alpha]
    return " ".join(tokens)


def extract_skills(text: str):
    if not text:
        return []
    text_lower = text.lower()
    return [s for s in SKILLS if s.lower() in text_lower]


def extract_keywords_advanced(text: str, top_n: int = 20):
    if not text or len(text.strip()) < 10:
        return []
    if nlp is None:
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        word_freq = Counter(words)
        return [w for w, _ in word_freq.most_common(top_n)]

    doc = nlp(text)
    keywords = [
        tok.lemma_.lower()
        for tok in doc
        if tok.pos_ in ["NOUN", "PROPN"] and not tok.is_stop and tok.is_alpha and len(tok.text) > 2
    ]
    entities = [ent.text.lower() for ent in doc.ents if ent.label_ in ['ORG', 'PRODUCT', 'TECHNOLOGY']]
    keywords.extend(entities)

    try:
        vectorizer = TfidfVectorizer(max_features=top_n, stop_words='english', ngram_range=(1, 2))
        tfidf_matrix = vectorizer.fit_transform([text])
        feature_names = vectorizer.get_feature_names_out()
        scores = tfidf_matrix.toarray()[0]
        tfidf_keywords = [feature_names[i] for i in scores.argsort()[-top_n:][::-1] if scores[i] > 0]
        all_keywords = list(set(keywords + tfidf_keywords))
        freq = Counter(all_keywords)
        return [w for w, _ in freq.most_common(top_n)]
    except Exception:
        freq = Counter(keywords)
        return [w for w, _ in freq.most_common(top_n)]


def extract_entities(text: str):
    if nlp is None or not text:
        return {}
    doc = nlp(text)
    entities = {}
    for ent in doc.ents:
        entities.setdefault(ent.label_, set()).add(ent.text)
    return {k: sorted(list(v)) for k, v in entities.items()}


def extract_contact_info(text: str):
    contact = {"email": [], "phone": [], "linkedin": [], "github": []}
    if not text:
        return contact

    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
    contact['email'] = sorted(list(set(emails)))

    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phones = re.findall(phone_pattern, text)
    contact['phone'] = sorted(list(set([''.join(p) for p in phones if any(p)])))

    linkedin = re.findall(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    contact['linkedin'] = sorted(list(set(linkedin)))

    github = re.findall(r'github\.com/[\w-]+', text, re.IGNORECASE)
    contact['github'] = sorted(list(set(github)))
    return contact


def load_resume_text(file_path: str) -> str:
    if file_path.lower().endswith('.pdf'):
        return extract_text_from_pdf_path(file_path)
    if file_path.lower().endswith('.docx'):
        return extract_text_from_docx_path(file_path)
    raise ValueError("Unsupported file format")


@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload():
    file = request.files.get('resume')
    top_n = int(request.form.get('top_keywords', 20))
    show_entities = request.form.get('show_entities') == 'on'
    show_contact = request.form.get('show_contact') == 'on'

    if not file or file.filename == '':
        flash('Please select a file to upload.', 'warning')
        return redirect(url_for('index'))

    try:
        if file.filename.lower().endswith('.pdf'):
            raw_text = extract_text_from_pdf_filelike(BytesIO(file.read()))
        elif file.filename.lower().endswith('.docx'):
            raw_text = extract_text_from_docx_filelike(BytesIO(file.read()))
        else:
            flash('Unsupported file format. Please upload a PDF or DOCX.', 'danger')
            return redirect(url_for('index'))

        keywords = extract_keywords_advanced(raw_text, top_n=top_n)
        skills = extract_skills(raw_text)
        entities = extract_entities(raw_text) if show_entities else {}
        contact = extract_contact_info(raw_text) if show_contact else {}

        # Keyword frequency for chart
        keyword_set = set([k.lower() for k in keywords])
        words = [w.strip().lower() for w in re.findall(r'\b[a-zA-Z]{3,}\b', raw_text)]
        freq = Counter([w for w in words if w in keyword_set])
        freq_data = [{"keyword": k, "count": v} for k, v in freq.most_common(15)]

        # Calculate summary statistics
        total_chars = len(raw_text)
        total_words = len(re.findall(r'\b\w+\b', raw_text))

        return render_template(
            'results.html',
            mode='single',
            filename=file.filename,
            raw_text=raw_text,
            keywords=keywords,
            skills=skills,
            entities=entities,
            contact=contact,
            keyword_freq=freq_data,
            total_chars=total_chars,
            total_words=total_words
        )
    except Exception as e:
        flash(f'Error processing file: {str(e)}', 'danger')
        return redirect(url_for('index'))


@app.route('/batch', methods=['POST'])
def batch():
    resume_folder = request.form.get('resume_folder', '').strip()
    job_description = request.form.get('job_description', '')
    top_n = int(request.form.get('top_keywords', 20))
    show_entities = request.form.get('show_entities') == 'on'
    show_contact = request.form.get('show_contact') == 'on'

    if not resume_folder or not os.path.exists(resume_folder):
        flash('Resume folder does not exist. Please provide a valid path.', 'warning')
        return redirect(url_for('index'))

    resume_files = [f for f in os.listdir(resume_folder) if f.lower().endswith(('.pdf', '.docx'))]
    if not resume_files:
        flash('No resume files found in the folder.', 'info')
        return redirect(url_for('index'))

    resume_texts = []
    processed = []
    for filename in resume_files:
        file_path = os.path.join(resume_folder, filename)
        try:
            raw_text = load_resume_text(file_path)
            clean_text = preprocess_text(raw_text)
            resume_texts.append(clean_text)

            item = {
                'filename': filename,
                'raw_text': raw_text,
                'keywords': extract_keywords_advanced(raw_text, top_n=top_n),
                'skills': extract_skills(raw_text),
                'entities': extract_entities(raw_text) if show_entities else {},
                'contact': extract_contact_info(raw_text) if show_contact else {}
            }
            processed.append(item)
        except Exception as e:
            # Skip files that fail and continue
            print(f"Error processing {filename}: {e}")
            continue

    if not processed:
        flash('Failed to process resumes.', 'danger')
        return redirect(url_for('index'))

    # Similarity against job description
    job_clean = preprocess_text(job_description)
    documents = resume_texts + [job_clean]
    try:
        vectorizer = TfidfVectorizer()
        tfidf = vectorizer.fit_transform(documents)
        job_vec = tfidf[-1]
        res_vecs = tfidf[:-1]
        sims = cosine_similarity(res_vecs, job_vec).flatten()
    except Exception:
        sims = [0.0] * len(processed)

    ranked = []
    # Match similarity scores with processed files (in correct order)
    # processed and sims are in the same order
    for info, score in zip(processed, sims):
        # Ensure score is between 0 and 1, then store as percentage
        normalized_score = max(0.0, min(1.0, float(score)))
        info['similarity_score'] = normalized_score
        ranked.append(info)

    ranked.sort(key=lambda x: x['similarity_score'], reverse=True)

    # Build keyword frequency across all
    all_kw = Counter()
    for r in ranked:
        all_kw.update([k.lower() for k in r['keywords']])
    keyword_freq = [{"keyword": k, "count": v} for k, v in all_kw.most_common(20)]

    return render_template(
        'results.html',
        mode='batch',
        ranked=ranked,
        keyword_freq=keyword_freq
    )


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8000, debug=True)
